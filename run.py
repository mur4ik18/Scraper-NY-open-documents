import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


# first link
GURL = "https://www.nysenate.gov/legislation/laws/CVP" 
# Find shell
Gshell = '#law-doc-wrapper > .c-law-link-container'
# Find title
Gtitle = 'a'
# All links for first page (Articles $)
firstPage_links = []
# All links for second pages(Articles $) > Sections $
secondPage_links = []

# Title(Article) 
contShell = '.node-statute'
# Title(Article) text 
Article = '.c-law--inactive-breadcrumb > span.crumb-label'
# Title(Article) numbers
Articles = '.c-law--inactive-breadcrumb > span.crumb-body'

# title
contTitle = '#doc-type'
# second title
contName = "#law-text-title"
# number sections
contLoc = '#location-id'
# text
contText = ".c-law-doc-text"



# docx document
document = Document()


# this function scrape first page
def fistPage(url, shell, mtitle):
    # get HTTP
    r = requests.get(url)
    # get HTML
    html = BeautifulSoup(r.content , 'html.parser')
    
    # links list
    lisst = []
    
    # find all lnks
    for el in html.select(shell):
        # find title and link
        title = el.select(mtitle)
        # get link end save in list
        lisst.append('https://www.nysenate.gov' + title[0].get('href'))
    # return list
    return lisst


# start page scrap and get links
firstPage_links = fistPage(GURL, Gshell, Gtitle)

# open every firstPage_links end scrap every link
for i in range(0 , len(firstPage_links)):
    # start page scrap
    secondPage_links.append(fistPage(firstPage_links[i], Gshell, Gtitle))


# this function scrape the second pages
def contentScrap(url,shell,mcontLoc, art, artTitle, contN, contT):
    # get HTTP
    r = requests.get(url)
    # get HTML
    html = BeautifulSoup(r.content , 'html.parser')
    # find all content 
    for el in html.select(shell):
        # subtitle 
        articleTitle = el.select(artTitle)
        # title num
        titlel = el.select(mcontLoc)
        # subtitle
        name = el.select(contN)
        # TEXT
        texts = el.select(contT)
        
        # get num
        text2 = titlel[0].text
        # get text subtitle
        text3 = name[0].text
        
        # print number + section
        print(text2.replace(' ', '')+ ' Section')
        
        # write header in docx
        d =document.add_heading(text2.replace(' ', '')+ ' Section', 2)
        # align center
        d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # write header in docx
        t=document.add_heading(text3.replace('  ', ''),3)
        # align center
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # free space
        document.add_paragraph('')
        # write text
        document.add_paragraph(texts[0].text)
        # free space
        document.add_paragraph('')

# find 'Article $' title
def fistPageTit(url, shell, mtitle):
    # get HTTP
    r = requests.get(url)
    # get HTML
    html = BeautifulSoup(r.content , 'html.parser')
    # title list
    lisst = []
    # find title
    for el in html.select(shell):
        # get title
        title = el.select(mtitle)
        # get title text
        lisst.append(title[0].text)
    # return title text
    return lisst




# 
titleList = []
# 
NameList = []
# get article name
titleList = fistPageTit('https://www.nysenate.gov/legislation/laws/CVP','.c-law-link-container', Gtitle)
# get article num
NameList = fistPageTit('https://www.nysenate.gov/legislation/laws/CVP','.c-law-link-title', 'a')

#
for i in range(0, len(secondPage_links)):
    # print Article
    print(titleList[i].replace('  ', ''))
    # print Article num
    print(NameList[i].replace('  ', ''))
    # write header
    p = document.add_heading(titleList[i].replace('  ', ''), 1)
    # align center
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # write header
    c =document.add_paragraph(NameList[i].replace('  ', ''))
    # align center
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # get content for every pages
    for j in range(0, len(secondPage_links[i])):        
        # start content scrap
        contentScrap(secondPage_links[i][j], contShell , contLoc, Article, Articles, contName, contText)


# close our document
document.save('done.docx')